from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, size, bold=True, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc = Document()

# Main Title
add_heading(doc, "LITTLE MILLENNIUM PRESCHOOL - PROJECT DOCUMENTATION", 16, center=True)
add_paragraph(doc, "")

# 1. Introduction
add_heading(doc, "1. Introduction", 14)
add_paragraph(doc, "The Little Millennium Preschool website has been fully upgraded from a static HTML/CSS structure to a modern, dynamic React application. This document outlines the technical details, architecture, and features of the newly built web platform to provide a clear understanding for all users, administrators, and developers.")
add_paragraph(doc, "")

# 2. Technology Stack
add_heading(doc, "2. Technology Stack", 14)
add_paragraph(doc, "The project utilizes the following modern web technologies:")
add_bullet(doc, "React.js: The core JavaScript library used to build interactive user interfaces.")
add_bullet(doc, "Vite: A fast build tool and development server that provides a quick and smooth development experience.")
add_bullet(doc, "React Router: A routing library used to navigate between different pages (Home, About, Admissions, etc.) without reloading the browser.")
add_bullet(doc, "CSS3: Standard CSS for styling, utilizing variables for consistent colors and fonts across the application.")
add_paragraph(doc, "")

# 3. Project Structure
add_heading(doc, "3. Project Structure", 14)
add_paragraph(doc, "The project files are organized cleanly to allow for easy maintenance:")
add_bullet(doc, "src/components: Contains shared layout elements like the Navbar, Footer, and the LoginModal.")
add_bullet(doc, "src/pages: Contains the individual pages (Home, About, Admissions, Classes, Teachers, Activities, Gallery, Contact).")
add_bullet(doc, "src/App.jsx: The main entry point that sets up the React Router and handles the overall layout.")
add_bullet(doc, "src/index.css: Contains the global styles, color variables, and typography definitions.")
add_paragraph(doc, "")

# 4. Key Rebranding Updates
add_heading(doc, "4. Key Rebranding Updates", 14)
add_paragraph(doc, "During the migration, the brand identity was updated across the entire website:")
add_bullet(doc, "Institution Name: Rebranded from 'Sunshine Kids' to 'Little Millennium Preschool'.")
add_bullet(doc, "Contact Number: Updated globally to +91 78923 26254.")
add_bullet(doc, "Address: Set to 222, 6th Cross Rd, Mico Layout, 2nd Stage Layout, BTM Layout, Bengaluru, Karnataka 560076.")
add_bullet(doc, "Operating Hours: Monday to Friday from 9 AM to 5 PM, Saturday from 9 AM to 12 PM, and Sunday Closed.")
add_paragraph(doc, "")

# 5. Core Features
add_heading(doc, "5. Core Features", 14)
add_bullet(doc, "Responsive Navigation: A fully responsive navigation bar with a mobile-friendly hamburger menu.")
add_bullet(doc, "Client-Side Routing: Instant page transitions without page reloads using React Router.")
add_bullet(doc, "Interactive Login Modal: A role-based login modal (Parent vs. Staff) driven by React state.")
add_bullet(doc, "Dynamic Content Display: Uses clean, state-driven components to display FAQs, gallery items, and class timetables.")
add_bullet(doc, "Contact Options: Integrated WhatsApp and Phone links for easy communication.")
add_paragraph(doc, "")

# 6. How to Run the Project
add_heading(doc, "6. How to Run the Project", 14)
add_paragraph(doc, "For developers and administrators looking to run or modify the code locally, follow these simple steps:")
add_bullet(doc, "1. Open the project folder 'sunshine-kids-react' in the terminal.")
add_bullet(doc, "2. Run 'npm install' to install all the necessary dependencies.")
add_bullet(doc, "3. Run 'npm run dev' to start the local development server.")
add_bullet(doc, "4. Open the provided localhost link in any modern web browser to view the application.")
add_paragraph(doc, "")

# 7. Conclusion
add_heading(doc, "7. Conclusion", 14)
add_paragraph(doc, "The Little Millennium Preschool website is now a highly interactive, scalable, and easy-to-maintain React application. It offers a smooth user experience for parents and a robust foundation for future enhancements.")

# Save the document
doc.save('Little_Millennium_Project_Documentation.docx')
